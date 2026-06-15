"""
End-to-end test for the template-aware DOCX export feature.

Flow:
  1. Register / login a test user
  2. Create a project
  3. Upload a content file (.txt)
  4. Upload a styled DOCX template (is_template=true) with custom fonts/margins
  5. Trigger report generation
  6. Poll until completed
  7. Download the DOCX
  8. Verify the output DOCX preserves the template's styles + has AI content
"""

import os, time, requests, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor

BASE = "http://127.0.0.1:8000/api/v1"
EMAIL = f"templatetest_{int(time.time())}@test.com"
PASSWORD = "TestPass123!"

def log(msg):
    # Safe print for Windows cp1252 terminal
    try:
        print(msg)
    except UnicodeEncodeError:
        print(msg.encode("ascii", "replace").decode("ascii"))


# ── Step 0: Create test fixtures ─────────────────────────────────────

# Content file
content_path = "test_fixture_content.txt"
with open(content_path, "w") as f:
    f.write(
        "Project Alpha is a machine learning pipeline for classifying satellite images.\n"
        "It uses a ResNet-50 backbone fine-tuned on 12,000 labeled images.\n"
        "Accuracy reached 94.7% on the holdout set.\n"
        "The pipeline runs on AWS SageMaker with a p3.2xlarge instance.\n"
    )

# DOCX template with distinctive styling
template_path = "test_fixture_template.docx"
doc = Document()

# Custom page margins
for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Customize Normal style
normal_style = doc.styles["Normal"]
normal_style.font.name = "Georgia"
normal_style.font.size = Pt(11)
normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Customize Heading 1
h1 = doc.styles["Heading 1"]
h1.font.name = "Georgia"
h1.font.size = Pt(24)
h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# Customize Heading 2
h2 = doc.styles["Heading 2"]
h2.font.name = "Georgia"
h2.font.size = Pt(16)
h2.font.color.rgb = RGBColor(0x2D, 0x2D, 0x5E)

# Add placeholder content (should be cleared by our export)
doc.add_heading("Company Report Template", level=1)
doc.add_paragraph("This is a placeholder paragraph that should be replaced.")
doc.add_heading("Section Placeholder", level=2)
doc.add_paragraph("Another placeholder.")
doc.save(template_path)
log(f"[OK] Created test fixtures: {content_path}, {template_path}")


# ── Step 1: Register + Login ─────────────────────────────────────────

r = requests.post(f"{BASE}/auth/register", json={
    "name": "Template Tester",
    "email": EMAIL,
    "password": PASSWORD
})
if r.status_code not in (200, 201):
    # Maybe user already exists, try login
    pass
log(f"[OK] Register: {r.status_code}")

r = requests.post(f"{BASE}/auth/login", json={"email": EMAIL, "password": PASSWORD})
assert r.status_code == 200, f"Login failed: {r.status_code} {r.text}"
token = r.json()["access_token"]
headers = {"Authorization": f"Bearer {token}"}
log(f"[OK] Logged in, token obtained")


# ── Step 2: Create a project ─────────────────────────────────────────

r = requests.post(f"{BASE}/projects", json={"project_name": "Template Test Project"}, headers=headers)
assert r.status_code in (200, 201), f"Create project failed: {r.status_code} {r.text}"
project_id = r.json()["id"]
log(f"[OK] Created project #{project_id}")


# ── Step 3: Upload content file ──────────────────────────────────────

with open(content_path, "rb") as f:
    r = requests.post(
        f"{BASE}/files/upload",
        data={"project_id": project_id, "is_template": "false"},
        files={"file": ("content.txt", f, "text/plain")},
        headers=headers
    )
assert r.status_code in (200, 201), f"Upload content failed: {r.status_code} {r.text}"
log(f"[OK] Uploaded content file")


# ── Step 4: Upload DOCX template ─────────────────────────────────────

with open(template_path, "rb") as f:
    r = requests.post(
        f"{BASE}/files/upload",
        data={"project_id": project_id, "is_template": "true"},
        files={"file": ("report_template.docx", f,
                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=headers
    )
assert r.status_code in (200, 201), f"Upload template failed: {r.status_code} {r.text}"
log(f"[OK] Uploaded DOCX template (is_template=true)")


# ── Step 5: Trigger report generation ────────────────────────────────

r = requests.post(f"{BASE}/reports", json={
    "project_id": project_id,
    "report_name": "Template E2E Test Report",
    "template_type": "custom"
}, headers=headers)
assert r.status_code in (200, 201), f"Create report failed: {r.status_code} {r.text}"
report_id = r.json()["id"]
log(f"[OK] Report #{report_id} created, generation started...")


# ── Step 6: Poll until done ──────────────────────────────────────────

MAX_WAIT = 300  # 5 minutes
start = time.time()
status = "pending"
while time.time() - start < MAX_WAIT:
    r = requests.get(f"{BASE}/reports/{report_id}", headers=headers)
    data = r.json()
    status = data["status"]
    elapsed = int(time.time() - start)
    log(f"  [{elapsed}s] Status: {status}")
    if status in ("completed", "failed"):
        break
    time.sleep(5)

if status == "failed":
    log(f"[FAIL] Report generation failed: {data.get('content', 'no details')}")
    sys.exit(1)
elif status != "completed":
    log(f"[FAIL] Timed out after {MAX_WAIT}s, last status: {status}")
    sys.exit(1)

log(f"[OK] Report completed in {int(time.time() - start)}s")


# ── Step 7: Download the DOCX ────────────────────────────────────────

r = requests.get(f"{BASE}/reports/{report_id}/download/docx", headers=headers)
assert r.status_code == 200, f"Download failed: {r.status_code} {r.text}"

download_path = "test_output_template_report.docx"
with open(download_path, "wb") as f:
    f.write(r.content)
log(f"[OK] Downloaded DOCX to {download_path} ({len(r.content)} bytes)")


# ── Step 8: Verify the output ────────────────────────────────────────

result = Document(download_path)

# 8a. Check page margins preserved
section = result.sections[0]
left_margin = round(section.left_margin.inches, 2)
top_margin = round(section.top_margin.inches, 2)
log(f"  Left margin: {left_margin} in (expected 1.25)")
log(f"  Top margin:  {top_margin} in (expected 1.50)")
assert left_margin == 1.25, f"Left margin mismatch: {left_margin}"
assert top_margin == 1.50, f"Top margin mismatch: {top_margin}"

# 8b. Check Normal style font preserved
normal = result.styles["Normal"]
log(f"  Normal font: {normal.font.name} (expected Georgia)")
log(f"  Normal size: {normal.font.size.pt if normal.font.size else 'inherited'}pt (expected 11)")
assert normal.font.name == "Georgia", f"Font mismatch: {normal.font.name}"
assert normal.font.size and normal.font.size.pt == 11.0, f"Size mismatch: {normal.font.size}"

# 8c. Check placeholder text was cleared
all_text = " ".join(p.text for p in result.paragraphs)
assert "placeholder" not in all_text.lower(), "Placeholder text was NOT cleared!"
log("  Placeholder text: cleared (good)")

# 8d. Check AI content was injected
assert len(result.paragraphs) > 3, f"Too few paragraphs: {len(result.paragraphs)}"
log(f"  Paragraph count: {len(result.paragraphs)}")

# Print first few paragraphs
log("\n  --- First 5 paragraphs of output ---")
for i, p in enumerate(result.paragraphs[:5]):
    if p.text.strip():
        log(f"  [{i}] ({p.style.name}): {p.text[:80]}")

log("\n========================================")
log("  ALL CHECKS PASSED!")
log("  Template styles preserved + AI content injected.")
log("========================================")

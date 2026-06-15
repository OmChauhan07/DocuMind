import os
import re
import shutil
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet


# ---------------------------------------------------------------------------
# Markdown → inline-run helper  (shared by both DOCX export paths)
# ---------------------------------------------------------------------------

def _add_markdown_runs(paragraph, text: str):
    """Parse **bold** and *italic* markers and add runs to a paragraph."""
    # Split on bold markers first
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _resolve_style(doc, preferred: str, fallback: str = "Normal"):
    """Return a style name that exists in the document, avoiding KeyError."""
    try:
        doc.styles[preferred]
        return preferred
    except KeyError:
        return fallback


# ---------------------------------------------------------------------------
# Path A – Generate a brand-new DOCX from markdown (no template)
# ---------------------------------------------------------------------------

def export_to_docx(content: str, output_path: str):
    """
    Exports a markdown string to a DOCX file.
    Provides basic parsing for headings, paragraphs, and bold text.
    """
    doc = Document()

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        p = None
        if line.startswith('### '):
            p = doc.add_heading(level=3)
            line = line[4:]
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            line = line[3:]
        elif line.startswith('# '):
            p = doc.add_heading(level=1)
            line = line[2:]
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            line = line[2:]
        else:
            p = doc.add_paragraph()

        _add_markdown_runs(p, line)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# Path B – Clone the user's template DOCX, preserve its formatting, inject
#           AI content using the template's own styles.
# ---------------------------------------------------------------------------

def export_to_docx_from_template(content: str, template_path: str, output_path: str):
    """
    Opens the user-provided template DOCX, clears its body text, and injects
    the AI-generated markdown content using the template's native styles so
    that fonts, margins, headers/footers, and page layout are preserved.
    """
    # 1. Copy the template to the output location so we never mutate the original
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copy2(template_path, output_path)

    # 2. Open the copy
    doc = Document(output_path)

    # 3. Clear existing body paragraphs (but keep section/page-setup/headers)
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Also clear any tables that were in the template body
    for tbl in list(doc.tables):
        tbl._element.getparent().remove(tbl._element)

    # 4. Discover the style names available in this template
    heading1_style = _resolve_style(doc, "Heading 1")
    heading2_style = _resolve_style(doc, "Heading 2")
    heading3_style = _resolve_style(doc, "Heading 3")
    bullet_style   = _resolve_style(doc, "List Bullet")
    normal_style   = _resolve_style(doc, "Normal")

    # 5. Inject AI content line-by-line, mapping markdown → template styles
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            # Preserve visual spacing
            doc.add_paragraph("", style=normal_style)
            continue

        if line.startswith('### '):
            p = doc.add_paragraph(style=heading3_style)
            _add_markdown_runs(p, line[4:])
        elif line.startswith('## '):
            p = doc.add_paragraph(style=heading2_style)
            _add_markdown_runs(p, line[3:])
        elif line.startswith('# '):
            p = doc.add_paragraph(style=heading1_style)
            _add_markdown_runs(p, line[2:])
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style=bullet_style)
            _add_markdown_runs(p, line[2:])
        elif line.startswith('|') and line.endswith('|'):
            # Simple markdown table row → just add as normal text for now
            p = doc.add_paragraph(style=normal_style)
            _add_markdown_runs(p, line)
        elif line.startswith('---'):
            # Horizontal rule → skip or add an empty line
            continue
        else:
            p = doc.add_paragraph(style=normal_style)
            _add_markdown_runs(p, line)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# PDF export (always generated fresh – no template concept for PDF)
# ---------------------------------------------------------------------------

def export_to_pdf(content: str, output_path: str):
    """
    Exports a markdown string to a PDF file using ReportLab.
    Provides basic parsing for headings, paragraphs, and bold text.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 12))
            continue

        # Replace **text** with <b>text</b> for ReportLab
        line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)

        if line.startswith('### '):
            story.append(Paragraph(line[4:], styles['Heading3']))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], styles['Heading2']))
        elif line.startswith('# '):
            story.append(Paragraph(line[2:], styles['Heading1']))
        elif line.startswith('- ') or line.startswith('* '):
            story.append(Paragraph(f"&bull; {line[2:]}", styles['Normal']))
        else:
            story.append(Paragraph(line, styles['Normal']))

    doc.build(story)
